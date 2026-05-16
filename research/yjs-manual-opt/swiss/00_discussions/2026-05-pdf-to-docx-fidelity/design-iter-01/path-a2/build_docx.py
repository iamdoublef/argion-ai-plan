"""
IMT050 PDF→可编辑 DOCX 复刻器（Path A2: python-docx 直接构建）

设计思路：
- 客户要的是可编辑 Word 文件用于局部修改，PDF 仅是视觉参考。
- 用 python-docx 从 JSON 内容直接构造 OOXML，所有文字都是真实段落和表格单元格。
- A5 页面，紧凑 spacing，10pt 字体，让 PDF 转换后 ≤16 页。
- 自渲染目录（不用 TOC field），用真实段落，制表符+前导点+页码。
- 章节标题：红色编号 + 黑色文本，左侧 4pt 黑色竖条（用段落左边框模拟）。
- 警告框 / 注意框 / 提示框 / 步骤序号：用表格 1×1 加底色 + 边框模拟。
- 图片：真实嵌入。
- Header/Footer：每个 Word section 一组，封面单独 section 无 H/F。

usage: python build_docx.py <out_docx>
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─── paths ────────────────────────────────────────────────────────────
HERE = Path(__file__).parent.resolve()
ROOT = Path("D:/work/private/yjsplan/research/yjs-manual-opt/swiss")
PRODUCT_DIR = ROOT / "products" / "imt050"
IMAGES_DIR = PRODUCT_DIR / "images"
CHAPTERS_DIR = PRODUCT_DIR / "content" / "source" / "chapters"
MANIFEST = PRODUCT_DIR / "content" / "source" / "manifest.json"
PRODUCT_JSON = PRODUCT_DIR / "product.json"
IMAGES_JSON = PRODUCT_DIR / "images.json"
STRINGS_ZHCN = PRODUCT_DIR / "i18n" / "compiled" / "zh-CN.json"

# ─── design tokens ────────────────────────────────────────────────────
PAGE_W_MM = 148.0
PAGE_H_MM = 210.0
PAGE_MARGIN_MM = 10.0

# Color tokens — matched to PDF target exactly
C_RED = RGBColor(0xE6, 0x39, 0x46)        # accent red
C_BLACK = RGBColor(0, 0, 0)
C_NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A) # primary text body
C_GRAY_BG = RGBColor(0xF4, 0xF4, 0xF4)    # table zebra (PDF uses #F4F4F4)
C_GRAY_TEXT = RGBColor(0x8E, 0x8E, 0x93)  # was 0x8A — PDF uses exactly #8E8E93
C_DARK_TEXT = RGBColor(0x22, 0x22, 0x22)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Fonts — match PDF target exactly
FONT_CN = "Microsoft YaHei"          # body CJK
FONT_CN_BOLD = "Microsoft YaHei"     # YaHei has built-in bold variant — set bold=True on the run
FONT_CN_SMALL = "NSimSun"            # small disclaimer (PDF uses NSimSun for tiny CJK)
FONT_LATIN = "Arial"                 # body Latin
FONT_LATIN_BLACK = "Arial Black"     # cover MODEL big number, chapter big number
FONT_MONO = "Courier New"            # CH.05 — SPECIFICATIONS header right, key chips

# ─── helpers ──────────────────────────────────────────────────────────
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
    """Each border: (sz_eighths_of_a_point, color_hex) or None."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tc_pr.append(tcBorders)
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        b = tcBorders.find(qn(f"w:{name}"))
        if b is not None:
            tcBorders.remove(b)
        b = OxmlElement(f"w:{name}")
        if val is None:
            b.set(qn("w:val"), "nil")
        else:
            sz, color = val
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), color)
        tcBorders.append(b)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80) -> None:
    """In twips (1/20 of a point). Default Word: top/bottom 0, left/right 108."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tc_pr.append(tcMar)


def set_paragraph_border(paragraph, *, left=None, top=None, bottom=None, right=None) -> None:
    """Each border: (sz, color_hex) or None."""
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = p_pr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        p_pr.append(pBdr)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        old = pBdr.find(qn(f"w:{name}"))
        if old is not None:
            pBdr.remove(old)
        if val is None:
            continue
        b = OxmlElement(f"w:{name}")
        sz, color = val
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)


def set_paragraph_shading(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def set_run_fonts(run, *, cn_font=FONT_CN, latin_font=FONT_LATIN, size_pt: Optional[float] = None,
                  bold: bool = False, color: Optional[RGBColor] = None, mono: bool = False,
                  latin_black: bool = False, small_cn: bool = False) -> None:
    """Apply font tokens to a run.

    Args:
        latin_black: use Arial Black for Latin (cover big text, chapter numbers).
        small_cn:    use NSimSun for CJK (disclaimer / very small text).
        mono:        use Courier New for both Latin and CJK (header right, MODEL, key chips).
    """
    if mono:
        cn_font = FONT_MONO
        latin_font = FONT_MONO
    if latin_black:
        latin_font = FONT_LATIN_BLACK
    if small_cn:
        cn_font = FONT_CN_SMALL
    run.font.name = latin_font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), latin_font)
    rFonts.set(qn("w:hAnsi"), latin_font)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:cs"), latin_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def add_run(p, text: str, **kw):
    r = p.add_run(text)
    set_run_fonts(r, **kw)
    return r


def parse_inline_md(paragraph, text: str, *, default_size: float = 7.0,
                    color: Optional[RGBColor] = None, mono_kw: bool = False) -> None:
    """Tiny markdown: **bold** segments. Newlines = soft line break."""
    if not text:
        return
    # Handle newlines first
    parts = text.split("\n")
    for li, line in enumerate(parts):
        if li > 0:
            br = OxmlElement("w:br")
            paragraph._p.add_element = None  # placeholder, we will add via run
        # split by **bold**
        segs = []
        i = 0
        cur = []
        while i < len(line):
            if line[i:i+2] == "**":
                # find closing **
                end = line.find("**", i+2)
                if end == -1:
                    cur.append(line[i:])
                    i = len(line)
                else:
                    if cur:
                        segs.append((False, "".join(cur)))
                        cur = []
                    segs.append((True, line[i+2:end]))
                    i = end + 2
            else:
                cur.append(line[i])
                i += 1
        if cur:
            segs.append((False, "".join(cur)))
        for bold, seg in segs:
            r = paragraph.add_run(seg)
            set_run_fonts(r, size_pt=default_size, bold=bold, color=color)
        if li < len(parts) - 1:
            # explicit line break
            run = paragraph.add_run("")
            set_run_fonts(run, size_pt=default_size, color=color)
            br = OxmlElement("w:br")
            run._r.append(br)


def set_paragraph_spacing(paragraph, *, before_pt: float = 0, after_pt: float = 0,
                          line_spacing: Optional[float] = None,
                          line_rule: str = "auto") -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        # if value <=3 treat as multiplier (1.15 etc), else as exact pt
        if line_rule == "exact":
            pPr = paragraph._p.get_or_add_pPr()
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:line"), str(int(line_spacing * 20)))
            spacing.set(qn("w:lineRule"), "exact")
        else:
            pf.line_spacing = line_spacing
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_page_size(section, w_mm: float, h_mm: float, margin_mm: float = 10.0,
                  header_mm: float = 12.0, footer_mm: float = 12.0) -> None:
    """Default page setup matched to PDF target:
      - margin_mm: outer page margin (10mm)
      - header_mm: header distance from top edge (was 4mm — bumped to 12mm to push header down)
      - footer_mm: footer distance from bottom edge (bumped to 12mm)
      - body top_margin is set to header_mm + a small gap so body starts below header
    """
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Mm(w_mm)
    section.page_height = Mm(h_mm)
    # body top margin: header_mm + a few mm gap below header line
    section.top_margin = Mm(header_mm + 7)
    section.bottom_margin = Mm(footer_mm + 5)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)
    section.header_distance = Mm(header_mm)
    section.footer_distance = Mm(footer_mm)


def make_clean_run(paragraph) -> Any:
    """Insert a transparent run so we can append OOXML elements."""
    r = paragraph.add_run("")
    return r


# ─── data loading ─────────────────────────────────────────────────────
def load_strings() -> Dict[str, str]:
    obj = json.loads(STRINGS_ZHCN.read_text(encoding="utf-8"))
    return obj["strings"]


def load_chapter(chapter_id: str) -> Dict[str, Any]:
    path = CHAPTERS_DIR / f"{chapter_id}.json"
    obj = json.loads(path.read_text(encoding="utf-8"))
    return obj


def load_manifest() -> Dict[str, Any]:
    return json.loads(MANIFEST.read_text(encoding="utf-8"))


def load_product() -> Dict[str, Any]:
    return json.loads(PRODUCT_JSON.read_text(encoding="utf-8"))


def load_images() -> Dict[str, Any]:
    return json.loads(IMAGES_JSON.read_text(encoding="utf-8"))


# ─── string lookup ────────────────────────────────────────────────────
class Strings:
    def __init__(self, table: Dict[str, str]):
        self.table = table

    def get(self, key: str, fallback: str = "") -> str:
        return self.table.get(key, fallback)

    def block_text(self, block: Dict[str, Any], fallback_field: str = "text") -> str:
        tid = block.get("text_id") or block.get("title_id")
        if tid and tid in self.table:
            return self.table[tid]
        return block.get(fallback_field) or block.get("title") or ""

    def item_text(self, item: Dict[str, Any]) -> str:
        tid = item.get("text_id")
        if tid and tid in self.table:
            return self.table[tid]
        return item.get("text", "")


# ─── builder ──────────────────────────────────────────────────────────
class ManualBuilder:
    def __init__(self, out_path: Path):
        self.out_path = out_path
        self.doc = Document()
        self.strings = Strings(load_strings())
        self.manifest = load_manifest()
        self.product = load_product()
        self.images = load_images()
        self._chapter_ref_current: str = ""
        self._page_no_per_chapter: Dict[str, int] = {}

    # ── core ────────────────────────────────────────────────────────
    def build(self):
        sec0 = self.doc.sections[0]
        set_page_size(sec0, PAGE_W_MM, PAGE_H_MM, PAGE_MARGIN_MM)
        sec0.different_first_page_header_footer = False  # cover section, no H/F at all

        # Wipe default Normal style spacing — set 7pt body to match PDF target
        normal = self.doc.styles["Normal"]
        normal.font.name = FONT_LATIN
        normal.font.size = Pt(7)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.35
        # Set eastAsia font globally on Normal style
        rPr = normal.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            normal.element.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), FONT_LATIN)
        rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:cs"), FONT_LATIN)

        # ── cover (section 1, no H/F)
        self._add_cover()

        # ── section break → TOC + body
        body_sec = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_size(body_sec, PAGE_W_MM, PAGE_H_MM, PAGE_MARGIN_MM)
        body_sec.header.is_linked_to_previous = False
        body_sec.footer.is_linked_to_previous = False
        self._setup_body_header_footer(body_sec)
        # restart page numbering at 2 for TOC; subsequent chapter sections continue
        self._set_page_start(body_sec, 2)

        # ── TOC page
        self._add_toc()

        # ── chapters
        for ch in self.manifest["chapters"]:
            if not ch.get("enabled", True):
                continue
            self._add_chapter(ch)

        # save
        self.doc.save(self.out_path)
        print(f"  wrote → {self.out_path}")

    # ── header/footer ──────────────────────────────────────────────
    def _set_page_start(self, section, start: int):
        """Set the section to restart page numbering at `start`."""
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType")
            sectPr.append(pgNumType)
        pgNumType.set(qn("w:start"), str(start))

    def _set_page_continue(self, section):
        """Ensure pgNumType has no 'start' attribute — page numbering continues from prev section."""
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is not None:
            # remove start attr if present
            if pgNumType.get(qn("w:start")) is not None:
                del pgNumType.attrib[qn("w:start")]

    def _setup_body_header_footer(self, section):
        """Header & footer matching PDF target:

        HEADER (top of every page):
          - Full-width black horizontal line (~1.5pt, sit on top)
          - left: 威富可 (YaHei Bold 7.5pt black)
          - right: CH.NN — SECTION (Courier New 5.25pt gray)
          - gap below header to body

        FOOTER (bottom):
          - gray "威富可  IMT050  说明书" (left, NSimSun 5pt gray)
          - page number (right, NSimSun 5pt gray)
        """
        header = section.header
        for p in list(header.paragraphs):
            p._p.getparent().remove(p._p)

        # Header paragraph with top border
        p = header.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=4, line_spacing=1.0)
        # top border (full width, thick black)
        set_paragraph_border(p, top=(12, "000000"))
        right_tab = Mm(PAGE_W_MM - 2 * PAGE_MARGIN_MM)
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(right_tab, alignment=WD_TAB_ALIGNMENT.RIGHT)
        # left: 威富可
        r1 = p.add_run("威富可")
        set_run_fonts(r1, size_pt=7.5, bold=True, color=C_BLACK)
        p.add_run("\t")
        # right: CH.XX — SECTION (Courier New, gray, 5.25pt)
        r2 = p.add_run("CHAPTER_REF_PLACEHOLDER")
        set_run_fonts(r2, size_pt=5.25, color=C_GRAY_TEXT, mono=True)

        # ── FOOTER
        footer = section.footer
        for p in list(footer.paragraphs):
            p._p.getparent().remove(p._p)
        p = footer.add_paragraph()
        set_paragraph_spacing(p, before_pt=4, after_pt=0, line_spacing=1.0)
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Mm(PAGE_W_MM - 2 * PAGE_MARGIN_MM),
                                  alignment=WD_TAB_ALIGNMENT.RIGHT)
        rf1 = p.add_run("威富可  IMT050  说明书")
        set_run_fonts(rf1, size_pt=5.5, color=C_GRAY_TEXT)
        p.add_run("\t")
        # PAGE field (page number — Arial Black for crisp look in PDF)
        run = p.add_run()
        set_run_fonts(run, size_pt=5.5, color=C_GRAY_TEXT, latin_black=True)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE   \\* MERGEFORMAT"
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)

    def _update_chapter_ref_header(self, section, chapter_ref: str):
        """Find and update the chapter-ref text in this section's header."""
        for p in section.header.paragraphs:
            for r in p.runs:
                if "CHAPTER_REF_PLACEHOLDER" in r.text or "CH." in r.text:
                    r.text = chapter_ref

    # ── cover page ─────────────────────────────────────────────────
    def _add_cover(self):
        """Cover layout — matches PDF target spans exactly:

        TARGET spans (from extract_spans.py):
          威富可        x=47.8  y=43.1   font=MicrosoftYaHei-Bold  size=7.5  color=#000000
          MODEL IMT050  x=28.3  y=328.8  font=CourierNewPS-Bold    size=6.0  color=#E63946
          制冰机        x=28.3  y=336.5  font=MicrosoftYaHei-Bold  size=18.0 color=#000000
          说明书        x=28.3  y=361.1  font=MicrosoftYaHei       size=7.5  color=#8E8E93
          disclaimer    NSimSun size=5.4 color=#8E8E93 (bottom)

        Page is 420x595pt (A5). Image is roughly 80-300y, narrow on left.
        Bottom area has a horizontal line + 2 lines of small disclaimer.
        """
        sec = self.doc.sections[0]
        sec.top_margin = Mm(14)
        sec.left_margin = Mm(10)
        sec.right_margin = Mm(10)
        sec.bottom_margin = Mm(26)  # room for footer disclaimer

        # ── Brand strip: short red bar + "威富可" (black)
        # PDF y=43.1 ≈ 15mm from top — that means short red bar is at very top
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=2, line_spacing=1.0)
        # red strip char + space + brand name
        r1 = p.add_run("━━  ")
        set_run_fonts(r1, size_pt=6.5, bold=True, color=C_RED)
        r2 = p.add_run("威富可")
        set_run_fonts(r2, size_pt=7.5, bold=True, color=C_BLACK)
        # very subtle inter-char spacing for elegance
        rPr = r2._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), "20")
        rPr.append(sp)

        # spacers to push image down to ~middle (page is ~210mm tall, image starts ~80mm)
        for _ in range(4):
            spc = self.doc.add_paragraph()
            set_paragraph_spacing(spc, before_pt=0, after_pt=0, line_spacing=1.0)
            spc.add_run("")

        # Product image (lower-left, ~30mm wide — PDF shows it small upper-mid)
        cover_path = IMAGES_DIR / self.product["product"]["cover_image"]
        if cover_path.exists():
            p_img = self.doc.add_paragraph()
            set_paragraph_spacing(p_img, before_pt=0, after_pt=0, line_spacing=1.0)
            p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p_img.add_run()
            run.add_picture(str(cover_path), width=Mm(30))  # ~PDF target proportion

        # ── MODEL IMT050 — 6pt, RED, Arial Black (small)
        p_model = self.doc.add_paragraph()
        set_paragraph_spacing(p_model, before_pt=4, after_pt=1, line_spacing=1.0)
        rm = p_model.add_run("M O D E L   I M T 0 5 0")
        set_run_fonts(rm, size_pt=6, bold=True, color=C_RED, mono=True)
        # tracking
        rPr = rm._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), "40")
        rPr.append(sp)

        # ── Title 制冰机 — 18pt, BLACK, YaHei Bold
        p_title = self.doc.add_paragraph()
        set_paragraph_spacing(p_title, before_pt=1, after_pt=2, line_spacing=1.0)
        rt = p_title.add_run("制冰机")
        set_run_fonts(rt, size_pt=18, bold=True, color=C_BLACK)

        # ── Subtitle 说明书 — 7.5pt, gray, YaHei
        p_sub = self.doc.add_paragraph()
        set_paragraph_spacing(p_sub, before_pt=2, after_pt=3, line_spacing=1.0)
        rs = p_sub.add_run("说明书")
        set_run_fonts(rs, size_pt=7.5, color=C_GRAY_TEXT)

        # ── Red short divider (under 说明书)
        p_div = self.doc.add_paragraph()
        set_paragraph_spacing(p_div, before_pt=1, after_pt=0, line_spacing=1.0)
        pf = p_div.paragraph_format
        pf.right_indent = Mm(PAGE_W_MM - 2*10 - 16)  # 16mm short bar
        set_paragraph_border(p_div, bottom=(18, "E63946"))
        r_div = p_div.add_run("")
        set_run_fonts(r_div, size_pt=2)

        # ── Cover footer: disclaimer with top hairline, NSimSun small gray
        cover_footer = sec.footer
        cover_footer.is_linked_to_previous = False
        for fp in list(cover_footer.paragraphs):
            fp._p.getparent().remove(fp._p)

        # Top border line above disclaimer
        p_top = cover_footer.add_paragraph()
        set_paragraph_spacing(p_top, before_pt=0, after_pt=2, line_spacing=1.0)
        set_paragraph_border(p_top, top=(6, "000000"))
        p_top.add_run("")

        # Disclaimer line 1 (smaller, gray, NSimSun)
        p_dis1 = cover_footer.add_paragraph()
        set_paragraph_spacing(p_dis1, before_pt=0, after_pt=0, line_spacing=1.25)
        rd1 = p_dis1.add_run("使用产品前请仔细阅读本说明书，并妥善保管。")
        set_run_fonts(rd1, size_pt=5.5, color=C_GRAY_TEXT, small_cn=True)
        # Disclaimer line 2
        p_dis2 = cover_footer.add_paragraph()
        set_paragraph_spacing(p_dis2, before_pt=0, after_pt=0, line_spacing=1.25)
        rd2 = p_dis2.add_run("说明书中的产品、配件等插图均为示意图，仅供参考。由于产品的更新与升级，产品实物与示意图可能略有差异，请以实物为准。")
        set_run_fonts(rd2, size_pt=5.5, color=C_GRAY_TEXT, small_cn=True)

    # ── TOC ────────────────────────────────────────────────────────
    def _add_toc(self):
        body_sec = self.doc.sections[1]
        self._update_chapter_ref_header(body_sec, "IMT050 — 说明书")

        # 目录 title — 18pt YaHei Bold, black (matches PDF chapter section title)
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=4, after_pt=10, line_spacing=1.05)
        rt = p.add_run("目录")
        set_run_fonts(rt, size_pt=18, bold=True, color=C_BLACK)

        target_pages = {
            "01-safety": 3, "02-usage-tips": 5, "03-structure": 6,
            "04-features": 7, "05-specifications": 8, "06-operation": 9,
            "07-troubleshooting": 11, "08-maintenance": 12,
            "09-installation-storage-disposal": 13, "10-warranty": 14,
        }
        for ch in self.manifest["chapters"]:
            if not ch.get("enabled", True):
                continue
            num = ch["chapter_no"]
            title = self.strings.get(ch["toc_title_id"], ch.get("toc_title", ""))
            page = target_pages.get(ch["chapter_id"], 1)
            self._add_toc_entry(num, title, page)

    def _add_toc_entry(self, num: str, title: str, page: int):
        """TOC row — PDF style:
          - red Arial Black '01' (7pt)
          - black YaHei Bold title (7pt)
          - gray page number on right (6pt)
          - no border below (cleaner than v05)
        """
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=4, after_pt=4, line_spacing=1.5)
        right_tab = Mm(PAGE_W_MM - 2 * PAGE_MARGIN_MM)
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(right_tab, alignment=WD_TAB_ALIGNMENT.RIGHT)

        # Arial Black '01' in red — matches PDF
        r_num = p.add_run(num)
        set_run_fonts(r_num, size_pt=7, bold=True, color=C_RED, latin_black=True)
        r_sep = p.add_run("    ")
        set_run_fonts(r_sep, size_pt=7)
        # YaHei Bold title in black
        r_title = p.add_run(title)
        set_run_fonts(r_title, size_pt=7.5, bold=True, color=C_BLACK)
        p.add_run("\t")
        # page no in gray
        r_page = p.add_run(str(page))
        set_run_fonts(r_page, size_pt=6.5, color=C_GRAY_TEXT)

    # ── chapter / page ─────────────────────────────────────────────
    def _add_chapter(self, manifest_chapter: Dict[str, Any]):
        cid = manifest_chapter["chapter_id"]
        chap = load_chapter(cid)
        header_ref = self.strings.get(manifest_chapter["header_ref_id"],
                                      manifest_chapter.get("header_ref", ""))

        # update header for this chapter — but Word headers are per-section,
        # not per-page. Strategy: each chapter becomes its own Word section
        # so the header chapter_ref can be different per chapter.
        # First chapter is already in body section 1 (TOC); we add new sections from 2nd.
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_size(new_section, PAGE_W_MM, PAGE_H_MM, PAGE_MARGIN_MM)
        # break link so each chapter section has its own header
        new_section.header.is_linked_to_previous = False
        new_section.footer.is_linked_to_previous = False
        self._setup_body_header_footer(new_section)
        self._set_page_continue(new_section)
        self._update_chapter_ref_header(new_section, header_ref)

        chapter_no = manifest_chapter["chapter_no"]
        chap_title = self.strings.get(manifest_chapter["title_id"],
                                      manifest_chapter.get("title", ""))

        # Render pages
        for pidx, page in enumerate(chap["pages"]):
            section_title = self.strings.get(page["section_title_id"], page.get("section_title", chap_title))
            if pidx > 0:
                self._page_break()
            self._add_section_title(chapter_no, section_title, pidx > 0)
            for block in page["blocks"]:
                self._render_block(block, chapter_no, page.get("page_class", ""))

    def _page_break(self):
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.0)
        r = p.add_run("")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r._r.append(br)

    def _add_section_title(self, chapter_no: str, title: str, is_continued: bool):
        """Chapter title — matches PDF spans exactly:
          chapter_no '05'  Arial-Black  13.5pt  #E63946
          title '技术参数' MicrosoftYaHei-Bold 13.5pt #000000
          + thick left black bar (≈3pt) full height
        """
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=4, line_spacing=1.0)
        # left bar: paragraph left border, 24 = 3pt
        set_paragraph_border(p, left=(24, "000000"))
        pf = p.paragraph_format
        pf.left_indent = Mm(2)
        # chapter num — Arial Black, red, 13.5pt
        r_num = p.add_run(chapter_no + " ")
        set_run_fonts(r_num, size_pt=13.5, bold=True, color=C_RED, latin_black=True)
        # title — YaHei Bold, black, 13.5pt
        r_t = p.add_run(title)
        set_run_fonts(r_t, size_pt=13.5, bold=True, color=C_BLACK)

    # ── blocks ─────────────────────────────────────────────────────
    def _render_block(self, block: Dict[str, Any], chapter_no: str, page_class: str):
        t = block.get("type")
        if t == "paragraph":
            self._render_paragraph(block)
        elif t == "sub_title":
            self._render_sub_title(block)
        elif t == "bullet_list":
            self._render_bullet_list(block)
        elif t == "warning_box":
            self._render_box(block, kind="warning")
        elif t == "caution_box":
            self._render_box(block, kind="caution")
        elif t == "notice_box":
            self._render_box(block, kind="notice")
        elif t == "step_flow":
            self._render_step_flow(block)
        elif t == "figure":
            self._render_figure(block)
        elif t == "figure_row":
            self._render_figure_row(block)
        elif t == "table_ref":
            self._render_table_ref(block)
        elif t == "warranty_card":
            self._render_warranty_card(block)
        else:
            print(f"  ! unknown block type: {t}")

    def _render_paragraph(self, block: Dict[str, Any]):
        """Body paragraph — 7pt YaHei (was 9pt) to match PDF density."""
        text = self.strings.block_text(block)
        text = self._substitute_vars(text)
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=1, after_pt=2, line_spacing=1.4)
        parse_inline_md(p, text, default_size=7, color=C_NEAR_BLACK)

    def _render_sub_title(self, block: Dict[str, Any]):
        """Sub-title (e.g. 放置与首次使用 / 用水要求) — bold black with thin gray underline.

        PDF shows sub-titles as bold black ~7pt followed by a thin gray horizontal rule.
        """
        text = self.strings.block_text(block)
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=4, after_pt=3, line_spacing=1.15)
        # Subtle gray bottom border for separator effect
        set_paragraph_border(p, bottom=(4, "E5E5EA"))
        r = p.add_run(text)
        set_run_fonts(r, size_pt=7.5, bold=True, color=C_BLACK)
        rPr = r._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), "10")
        rPr.append(sp)

    def _render_bullet_list(self, block: Dict[str, Any]):
        """Bullet list — PDF uses red asterisk (*) prefix at body size."""
        for item in block.get("items", []):
            text = self.strings.item_text(item)
            if not text:
                continue
            text = self._substitute_vars(text)
            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.35)
            pf = p.paragraph_format
            pf.left_indent = Mm(4.5)
            pf.first_line_indent = Mm(-3)
            r0 = p.add_run("*  ")
            set_run_fonts(r0, size_pt=8, bold=True, color=C_RED)
            parse_inline_md(p, text, default_size=7, color=C_NEAR_BLACK)

    def _render_box(self, block: Dict[str, Any], *, kind: str):
        """Warning/Caution/Notice box — matches PDF design:

        - 1x1 table full width
        - Warning: red border, red ▲ icon, "▲ 警告 WARNING" title row
        - Caution: black border, black title
        - Notice: thin gray border + light gray shade
        - body bullets 7pt YaHei
        """
        title_id = block.get("title_id")
        if title_id and title_id in self.strings.table:
            title = self.strings.table[title_id]
        else:
            title = block.get("title", "")
        items = block.get("items", [])
        text_id = block.get("text_id")
        if text_id and text_id in self.strings.table:
            text = self.strings.table[text_id]
        else:
            text = block.get("text", "")

        # Per-kind styling
        if kind == "warning":
            border_color = "E63946"
            shading = "FFFFFF"
            border_width = 8
            icon_char = "▲"
            icon_color = C_RED
        elif kind == "caution":
            border_color = "000000"
            shading = "FFFFFF"
            border_width = 6
            icon_char = "▲"
            icon_color = C_BLACK
        else:  # notice
            border_color = "8E8E93"
            shading = "F4F4F4"
            border_width = 4
            icon_char = ""
            icon_color = C_BLACK

        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False
        tbl.columns[0].width = Mm(usable_w)
        cell = tbl.cell(0, 0)
        cell.width = Mm(usable_w)
        set_cell_shading(cell, shading)
        set_cell_borders(cell,
                         top=(border_width, border_color),
                         bottom=(border_width, border_color),
                         left=(border_width, border_color),
                         right=(border_width, border_color))
        set_cell_margins(cell, top=70, bottom=70, left=100, right=100)

        # Title row: icon + title
        first_para = cell.paragraphs[0]
        set_paragraph_spacing(first_para, before_pt=0, after_pt=3, line_spacing=1.0)
        if icon_char:
            icon_run = first_para.add_run(icon_char + " ")
            set_run_fonts(icon_run, size_pt=9, bold=True, color=icon_color)
        rt = first_para.add_run(title)
        set_run_fonts(rt, size_pt=7.5, bold=True, color=C_BLACK)

        # items as list
        if items:
            for item in items:
                t = self.strings.item_text(item)
                if not t:
                    continue
                p = cell.add_paragraph()
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.35)
                pf = p.paragraph_format
                pf.left_indent = Mm(4)
                pf.first_line_indent = Mm(-3)
                r0 = p.add_run("*  ")
                set_run_fonts(r0, size_pt=7, bold=True, color=C_RED)
                parse_inline_md(p, t, default_size=6.5, color=C_NEAR_BLACK)

        # body text if no items
        if text and not items:
            p = cell.add_paragraph()
            set_paragraph_spacing(p, before_pt=1, after_pt=0, line_spacing=1.35)
            parse_inline_md(p, text, default_size=6.5, color=C_NEAR_BLACK)

        # tiny spacer after box
        sp = self.doc.add_paragraph()
        set_paragraph_spacing(sp, before_pt=0, after_pt=1, line_spacing=1.0)
        sp.add_run("")

    def _render_step_flow(self, block: Dict[str, Any]):
        """Numbered step flow — PDF style:
          - small black square with white Arial Black number (7pt)
          - body text 7pt YaHei
        """
        start_at = int(block.get("start_at", 1))
        steps = block.get("steps", [])
        default_max_h = self._mm_value(block.get("max_height"), default=24)
        for idx, step in enumerate(steps):
            n = start_at + idx
            text = self.strings.item_text(step) if "id" in step else (
                self.strings.get(step.get("text_id", ""), step.get("text", "")))
            tbl = self.doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
            num_col_w = 4.5
            text_col_w = usable_w - num_col_w
            tbl.columns[0].width = Mm(num_col_w)
            tbl.columns[1].width = Mm(text_col_w)
            tbl.rows[0].height = Mm(4.5)
            tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            # num cell — small BLACK square with WHITE Arial Black number (matches PDF)
            c0 = tbl.cell(0, 0)
            c0.width = Mm(num_col_w)
            set_cell_shading(c0, "1A1A1A")
            set_cell_borders(c0,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=(8, "1A1A1A"), right=(8, "1A1A1A"))
            set_cell_margins(c0, top=10, bottom=10, left=20, right=20)
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p0, before_pt=0, after_pt=0, line_spacing=1.0)
            rn = p0.add_run(str(n))
            set_run_fonts(rn, size_pt=7, bold=True, color=C_WHITE, latin_black=True)
            # text cell
            c1 = tbl.cell(0, 1)
            c1.width = Mm(text_col_w)
            set_cell_margins(c1, top=10, bottom=10, left=100, right=20)
            c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p1 = c1.paragraphs[0]
            set_paragraph_spacing(p1, before_pt=0, after_pt=0, line_spacing=1.35)
            parse_inline_md(p1, text, default_size=7, color=C_NEAR_BLACK)
            # figures (if any)
            figs = step.get("figures") or []
            if figs:
                step_max_h = self._mm_value(step.get("max_height"), default=default_max_h)
                # If multiple figures in same step, lay them side by side in one row
                if len(figs) == 1:
                    self._add_image_paragraph(figs[0], max_height_mm=step_max_h, indent_mm=8, center=False)
                else:
                    self._add_image_row(figs, max_height_mm=step_max_h)
            # spacer (very small)
            sp = self.doc.add_paragraph()
            set_paragraph_spacing(sp, before_pt=0, after_pt=0, line_spacing=1.0)
            sp.add_run("")

    def _add_image_row(self, fig_ids: List[str], *, max_height_mm: float = 24):
        n = len(fig_ids)
        if n == 0:
            return
        tbl = self.doc.add_table(rows=1, cols=n)
        tbl.autofit = False
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        col_w = usable_w / n
        # Reserve some inner margin
        img_cap_w = col_w * 0.92
        for ci, fid in enumerate(fig_ids):
            tbl.columns[ci].width = Mm(col_w)
            cell = tbl.cell(0, ci)
            cell.width = Mm(col_w)
            set_cell_margins(cell, top=10, bottom=10, left=20, right=20)
            set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_path = self._image_path(fid)
            if img_path and img_path.exists():
                run = p.add_run()
                try:
                    # pick width if image is wide, height otherwise
                    try:
                        from PIL import Image as _PIL
                        with _PIL.open(img_path) as im:
                            w_px, h_px = im.size
                    except Exception:
                        w_px, h_px = (1, 1)
                    if w_px / max(h_px, 1) > img_cap_w / max_height_mm:
                        run.add_picture(str(img_path), width=Mm(img_cap_w))
                    else:
                        run.add_picture(str(img_path), height=Mm(max_height_mm))
                except Exception:
                    pass

    def _render_figure(self, block: Dict[str, Any]):
        fig_id = block.get("figure")
        if not fig_id:
            return
        usable_w_mm = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        # Determine image aspect ratio
        img_path = self._image_path(fig_id)
        aspect = 1.0
        if img_path and img_path.exists():
            try:
                from PIL import Image as _PILImage
                with _PILImage.open(img_path) as im:
                    w_px, h_px = im.size
                aspect = w_px / h_px if h_px else 1.0
            except Exception:
                pass
        if block.get("max_height"):
            max_h = self._mm_value(block.get("max_height"), default=40)
            # Behave like HTML: image natural width up to ~85% page-width, capped by max-height
            max_w = usable_w_mm * 0.85
            self._add_image_paragraph(fig_id, max_height_mm=max_h, max_width_mm=max_w, center=True)
            return
        if aspect > 6:
            # very wide → decorative separator. fit to width, ~5mm tall max
            self._add_image_paragraph(fig_id, max_height_mm=5, max_width_mm=usable_w_mm * 0.85, center=True)
        elif aspect > 2:
            self._add_image_paragraph(fig_id, max_height_mm=18, max_width_mm=usable_w_mm * 0.85, center=True)
        else:
            self._add_image_paragraph(fig_id, max_height_mm=30, max_width_mm=usable_w_mm * 0.85, center=True)

    def _render_figure_row(self, block: Dict[str, Any]):
        figs = block.get("figures") or []
        items = block.get("items") or []
        if figs:
            tbl = self.doc.add_table(rows=1, cols=len(figs))
            tbl.autofit = False
            usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
            col_w = usable_w / len(figs)
            for ci, fig in enumerate(figs):
                tbl.columns[ci].width = Mm(col_w)
                cell = tbl.cell(0, ci)
                cell.width = Mm(col_w)
                set_cell_margins(cell, top=20, bottom=20, left=40, right=40)
                set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
                p = cell.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_id = fig.get("figure")
                max_h_mm = self._mm_value(fig.get("max_height"), default=28)
                max_w_pct = fig.get("max_width", "100%")
                max_w_mm = col_w
                if isinstance(max_w_pct, str) and max_w_pct.endswith("%"):
                    pct = float(max_w_pct[:-1]) / 100.0
                    max_w_mm = col_w * pct
                img_path = self._image_path(fig_id)
                if img_path and img_path.exists():
                    run = p.add_run()
                    try:
                        run.add_picture(str(img_path), height=Mm(max_h_mm))
                    except Exception:
                        pass
        elif items:
            # status-indicator-row: [label_before][image]
            tbl = self.doc.add_table(rows=1, cols=len(items)*2)
            tbl.autofit = False
            usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
            col_w = usable_w / (len(items)*2)
            for i, item in enumerate(items):
                # label cell
                cell_l = tbl.cell(0, i*2)
                cell_l.width = Mm(col_w)
                set_cell_margins(cell_l, top=20, bottom=20, left=40, right=40)
                set_cell_borders(cell_l, top=None, bottom=None, left=None, right=None)
                p = cell_l.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                label = self.strings.get(item.get("label_before_id", ""), item.get("label_before", ""))
                r = p.add_run(label)
                set_run_fonts(r, size_pt=7, bold=True, color=C_NEAR_BLACK)
                # image cell
                cell_i = tbl.cell(0, i*2+1)
                cell_i.width = Mm(col_w)
                set_cell_margins(cell_i, top=20, bottom=20, left=40, right=40)
                set_cell_borders(cell_i, top=None, bottom=None, left=None, right=None)
                pi = cell_i.paragraphs[0]
                set_paragraph_spacing(pi, before_pt=0, after_pt=0, line_spacing=1.0)
                pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell_i.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                fig_id = item.get("figure")
                img_path = self._image_path(fig_id)
                if img_path and img_path.exists():
                    run = pi.add_run()
                    try:
                        run.add_picture(str(img_path), height=Mm(4))
                    except Exception:
                        pass

    def _render_table_ref(self, block: Dict[str, Any]):
        tname = block.get("table")
        if tname == "specs":
            self._render_specs_table()
        elif tname == "parts":
            self._render_parts_table()
        elif tname == "buttons":
            self._render_buttons_table()
        elif tname == "brand_info":
            self._render_brand_info_table()
        elif tname == "manufacturer_info":
            self._render_manufacturer_info_table()
        elif "headers" in block and "rows" in block:
            # inline table_ref (troubleshooting)
            self._render_inline_table(block)
        else:
            print(f"  ! unknown table: {tname}")

    def _render_inline_table(self, block: Dict[str, Any]):
        """Generic inline table (troubleshooting). PDF-style:
          - HEADER row: BLACK fill, WHITE bold text
          - Body rows: only top/bottom thin gray borders, no left/right
          - No zebra by default (PDF doesn't use it on this style)
        """
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        # parse widths in % (default share leftover)
        pct_sum = 0.0
        widths_mm = []
        for h in headers:
            w = h.get("width", "")
            if isinstance(w, str) and w.endswith("%"):
                pp = float(w[:-1])
                widths_mm.append(usable_w * pp / 100.0)
                pct_sum += pp
            else:
                widths_mm.append(None)
        leftover_pct = max(0, 100 - pct_sum)
        unspecified = [i for i, w in enumerate(widths_mm) if w is None]
        if unspecified:
            share = (usable_w * leftover_pct / 100.0) / len(unspecified)
            for i in unspecified:
                widths_mm[i] = share

        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.autofit = False
        for ci, w in enumerate(widths_mm):
            tbl.columns[ci].width = Mm(w)

        # HEADER row — BLACK background, WHITE bold 6pt
        for ci, h in enumerate(headers):
            cell = tbl.cell(0, ci)
            cell.width = Mm(widths_mm[ci])
            set_cell_shading(cell, "1A1A1A")
            set_cell_borders(cell,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=None, right=None)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
            txt = self.strings.get(h.get("text_id", ""), h.get("text", ""))
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.1)
            r = p.add_run(txt)
            set_run_fonts(r, size_pt=7, bold=True, color=C_WHITE)
        # body rows — only horizontal lines, 7pt body
        for ri, row in enumerate(rows):
            for ci, cell_data in enumerate(row):
                cell = tbl.cell(ri+1, ci)
                cell.width = Mm(widths_mm[ci])
                set_cell_borders(cell,
                                 top=None,
                                 bottom=(4, "E5E5EA"),
                                 left=None, right=None)
                set_cell_margins(cell, top=30, bottom=30, left=80, right=60)
                txt = self.strings.get(cell_data.get("text_id", ""), cell_data.get("text", ""))
                p = cell.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.3)
                parse_inline_md(p, txt, default_size=7, color=C_NEAR_BLACK)

        # spacer
        sp = self.doc.add_paragraph()
        set_paragraph_spacing(sp, before_pt=0, after_pt=2, line_spacing=1.0)
        sp.add_run("")

    def _render_specs_table(self):
        """Specs table — matches PDF exactly:
          - HEADER: BLACK fill, WHITE bold 6pt
          - body: 6.5pt YaHei/Arial, only thin gray bottom borders
        """
        ui = self.product["ui_labels"]["cn"]
        market = self.product["product"].get("active_market", "eu")
        specs_rows = self.product["specs"][market]["rows"]
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        col1_w = usable_w * 0.42
        col2_w = usable_w - col1_w
        tbl = self.doc.add_table(rows=1 + len(specs_rows), cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Mm(col1_w)
        tbl.columns[1].width = Mm(col2_w)
        # HEADER row — BLACK fill, WHITE bold 6pt
        for ci, label in enumerate([ui["specs_col1"], ui["specs_col2"]]):
            cell = tbl.cell(0, ci)
            cell.width = Mm(col1_w if ci == 0 else col2_w)
            set_cell_shading(cell, "1A1A1A")
            set_cell_borders(cell,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=None, right=None)
            set_cell_margins(cell, top=30, bottom=30, left=80, right=70)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.1)
            r = p.add_run(label)
            set_run_fonts(r, size_pt=7, bold=True, color=C_WHITE)
        # body rows — only thin gray bottom border, 7pt body (matches target 6.67pt rendering well)
        for ri, row in enumerate(specs_rows):
            cell0 = tbl.cell(ri+1, 0)
            cell1 = tbl.cell(ri+1, 1)
            for c in (cell0, cell1):
                c.width = Mm(col1_w if c is cell0 else col2_w)
                set_cell_borders(c,
                                 top=None,
                                 bottom=(4, "E5E5EA"),
                                 left=None, right=None)
                set_cell_margins(c, top=30, bottom=30, left=80, right=70)
                p = c.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.25)
            r0 = cell0.paragraphs[0].add_run(row["label"])
            set_run_fonts(r0, size_pt=7, color=C_NEAR_BLACK)
            r1 = cell1.paragraphs[0].add_run(row["value"])
            set_run_fonts(r1, size_pt=7, color=C_NEAR_BLACK)

    def _render_parts_table(self):
        """Parts table — PDF-style:
          - BLACK header row [编号|名称|编号|名称]
          - body: only horizontal lines
          - 7pt body text
        """
        parts = self.product["parts"]
        ui = self.product["ui_labels"]["cn"]
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        c1_w = usable_w * 0.10
        c2_w = usable_w * 0.40
        c3_w = usable_w * 0.10
        c4_w = usable_w * 0.40
        n = len(parts)
        rows = (n + 1) // 2
        tbl = self.doc.add_table(rows=1 + rows, cols=4)
        tbl.autofit = False
        widths = [c1_w, c2_w, c3_w, c4_w]
        for ci, w in enumerate(widths):
            tbl.columns[ci].width = Mm(w)
        # HEADER row — BLACK fill, WHITE 6pt bold
        header_labels = [ui["parts_col_no"], ui["parts_col_name"],
                         ui["parts_col_no"], ui["parts_col_name"]]
        for ci, lbl in enumerate(header_labels):
            cell = tbl.cell(0, ci)
            cell.width = Mm(widths[ci])
            set_cell_shading(cell, "1A1A1A")
            set_cell_borders(cell,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=None, right=None)
            set_cell_margins(cell, top=30, bottom=30, left=70, right=60)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.1)
            r = p.add_run(lbl)
            set_run_fonts(r, size_pt=7, bold=True, color=C_WHITE)
        # body — only horizontal lines, 6.5pt
        for ri in range(rows):
            left = parts[ri]
            right = parts[ri + rows] if ri + rows < n else None
            entries = [("id", left), ("name", left)]
            if right:
                entries += [("id", right), ("name", right)]
            else:
                entries += [(None, None), (None, None)]
            for ci, (kind, part) in enumerate(entries):
                cell = tbl.cell(ri+1, ci)
                cell.width = Mm(widths[ci])
                set_cell_borders(cell,
                                 top=None,
                                 bottom=(4, "E5E5EA"),
                                 left=None, right=None)
                set_cell_margins(cell, top=30, bottom=30, left=70, right=60)
                p = cell.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.3)
                if part is None:
                    continue
                if kind == "id":
                    r = p.add_run(str(part["id"]))
                    set_run_fonts(r, size_pt=7, color=C_NEAR_BLACK)
                else:
                    r = p.add_run(part["name_cn"])
                    set_run_fonts(r, size_pt=7, color=C_NEAR_BLACK)

    def _render_buttons_table(self):
        """Buttons table — PDF-style:
          - BLACK header
          - body: only horizontal lines
        """
        ui = self.product["ui_labels"]["cn"]
        buttons = self.product["buttons"]
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        c1_w = usable_w * 0.30
        c2_w = usable_w - c1_w
        tbl = self.doc.add_table(rows=1 + len(buttons), cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Mm(c1_w)
        tbl.columns[1].width = Mm(c2_w)
        # HEADER — BLACK
        for ci, lbl in enumerate([ui["buttons_col1"], ui["buttons_col2"]]):
            cell = tbl.cell(0, ci)
            cell.width = Mm(c1_w if ci == 0 else c2_w)
            set_cell_shading(cell, "1A1A1A")
            set_cell_borders(cell,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=None, right=None)
            set_cell_margins(cell, top=30, bottom=30, left=70, right=60)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.1)
            r = p.add_run(lbl)
            set_run_fonts(r, size_pt=7, bold=True, color=C_WHITE)
        # body — only horizontal lines
        for ri, b in enumerate(buttons):
            cell0 = tbl.cell(ri+1, 0)
            cell1 = tbl.cell(ri+1, 1)
            for c, ww in ((cell0, c1_w), (cell1, c2_w)):
                c.width = Mm(ww)
                set_cell_borders(c,
                                 top=None,
                                 bottom=(4, "E5E5EA"),
                                 left=None, right=None)
                set_cell_margins(c, top=40, bottom=40, left=70, right=60)
                p = c.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.3)
            p0 = cell0.paragraphs[0]
            r_n = p0.add_run(f"{b['id']} ")
            set_run_fonts(r_n, size_pt=7, bold=True, color=C_RED, latin_black=True)
            # key chip (monospace, inside [])
            r_key = p0.add_run(b["key"])
            set_run_fonts(r_key, size_pt=6.5, bold=True, color=C_BLACK, mono=True)
            r_sp = p0.add_run("   ")
            set_run_fonts(r_sp, size_pt=7)
            r_name = p0.add_run(b["name_cn"])
            set_run_fonts(r_name, size_pt=7, color=C_NEAR_BLACK)
            p1 = cell1.paragraphs[0]
            r_desc = p1.add_run(b["desc_cn"])
            set_run_fonts(r_desc, size_pt=7, color=C_NEAR_BLACK)

    def _render_brand_info_table(self):
        ui = self.product["ui_labels"]["cn"]
        brand = self.product["brands"]["wevac"]
        rows = [
            (ui["brand_label"], brand["name"], True),
            (ui["address_label"], brand["address_cn"], False),
            (ui["website_label"], brand["website"], False),
            (ui["email_label"], brand["support_email"], False),
        ]
        self._render_two_col_kv_table(rows)

    def _render_manufacturer_info_table(self):
        ui = self.product["ui_labels"]["cn"]
        mfg = self.product["manufacturer"]
        rows = [
            (ui["manufacturer_label"], f"{mfg['name_cn']}\n{mfg['name_en']}", True),
            (ui["address_label"], mfg["address_cn"], False),
            (ui["website_label"], mfg["website"], False),
        ]
        self._render_two_col_kv_table(rows)

    def _render_two_col_kv_table(self, rows: List[tuple]):
        """KV table for brand/manufacturer info — PDF-style:
          - BLACK header row [项目 | 信息]
          - body: only horizontal lines
        """
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        col1_w = usable_w * 0.30
        col2_w = usable_w - col1_w
        # +1 row for header
        tbl = self.doc.add_table(rows=len(rows) + 1, cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Mm(col1_w)
        tbl.columns[1].width = Mm(col2_w)
        # HEADER row
        header_labels = ["项目", "信息"]
        for ci, lbl in enumerate(header_labels):
            cell = tbl.cell(0, ci)
            cell.width = Mm(col1_w if ci == 0 else col2_w)
            set_cell_shading(cell, "1A1A1A")
            set_cell_borders(cell,
                             top=(8, "1A1A1A"), bottom=(8, "1A1A1A"),
                             left=None, right=None)
            set_cell_margins(cell, top=30, bottom=30, left=70, right=60)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.1)
            r = p.add_run(lbl)
            set_run_fonts(r, size_pt=7, bold=True, color=C_WHITE)
        for ri, (label, value, bold) in enumerate(rows):
            cell0 = tbl.cell(ri + 1, 0)
            cell1 = tbl.cell(ri + 1, 1)
            for c, ww in ((cell0, col1_w), (cell1, col2_w)):
                c.width = Mm(ww)
                set_cell_borders(c,
                                 top=None,
                                 bottom=(4, "E5E5EA"),
                                 left=None, right=None)
                set_cell_margins(c, top=40, bottom=40, left=70, right=60)
                p = c.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.35)
            r0 = cell0.paragraphs[0].add_run(label)
            set_run_fonts(r0, size_pt=7, color=C_NEAR_BLACK)
            parse_inline_md(cell1.paragraphs[0], value, default_size=7, color=C_NEAR_BLACK)
            if bold:
                for run in cell1.paragraphs[0].runs:
                    run.font.bold = True

    def _render_warranty_card(self, block: Dict[str, Any]):
        """Warranty card — PDF uses ZEBRA stripes here, full grid border to be 'filled in'."""
        fields = block.get("fields", [])
        usable_w = PAGE_W_MM - 2 * PAGE_MARGIN_MM
        col1_w = usable_w * 0.35
        col2_w = usable_w - col1_w
        tbl = self.doc.add_table(rows=len(fields), cols=2)
        tbl.autofit = False
        tbl.columns[0].width = Mm(col1_w)
        tbl.columns[1].width = Mm(col2_w)
        for ri, field in enumerate(fields):
            label = self.strings.get(field.get("text_id", ""), field.get("text", ""))
            cell0 = tbl.cell(ri, 0)
            cell1 = tbl.cell(ri, 1)
            for c, ww in ((cell0, col1_w), (cell1, col2_w)):
                c.width = Mm(ww)
                set_cell_borders(c,
                                 top=(4, "C0C0C0"), bottom=(4, "C0C0C0"),
                                 left=(4, "C0C0C0"), right=(4, "C0C0C0"))
                set_cell_margins(c, top=80, bottom=80, left=80, right=80)
                # zebra stripes — PDF uses #F4F4F4 for even rows
                if ri % 2 == 0:
                    set_cell_shading(c, "F4F4F4")
                p = c.paragraphs[0]
                set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=1.3)
            r = cell0.paragraphs[0].add_run(label)
            set_run_fonts(r, size_pt=7, color=C_NEAR_BLACK)
            cell1.paragraphs[0].add_run("")

    # ── helpers for images & sizes ─────────────────────────────────
    def _image_path(self, fig_id: Optional[str]) -> Optional[Path]:
        if not fig_id:
            return None
        info = self.images.get(fig_id)
        if not info:
            return None
        return IMAGES_DIR / info["file"]

    def _add_image_paragraph(self, fig_id: str, *, max_height_mm: float = 30,
                             max_width_mm: Optional[float] = None,
                             center: bool = True, indent_mm: float = 0):
        path = self._image_path(fig_id)
        if path is None or not path.exists():
            return
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=2, after_pt=2, line_spacing=1.0)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif indent_mm:
            p.paragraph_format.left_indent = Mm(indent_mm)
        run = p.add_run()
        try:
            # Inspect aspect ratio and apply the tighter bound
            try:
                from PIL import Image as _PILImage
                with _PILImage.open(path) as im:
                    w_px, h_px = im.size
            except Exception:
                w_px, h_px = (1, 1)
            if max_width_mm is not None and max_height_mm is not None and w_px > 0 and h_px > 0:
                # both — pick whichever yields smaller image
                if w_px / h_px > max_width_mm / max_height_mm:
                    # width-limited
                    run.add_picture(str(path), width=Mm(max_width_mm))
                else:
                    run.add_picture(str(path), height=Mm(max_height_mm))
            elif max_width_mm is not None:
                run.add_picture(str(path), width=Mm(max_width_mm))
            else:
                run.add_picture(str(path), height=Mm(max_height_mm))
        except Exception as e:
            print(f"  ! image {fig_id} failed: {e}")

    def _mm_value(self, css_size: Optional[str], default: float = 30.0) -> float:
        if not css_size:
            return default
        s = str(css_size).strip()
        if s.endswith("mm"):
            return float(s[:-2])
        if s.endswith("px"):
            return float(s[:-2]) * 25.4 / 96.0
        if s.endswith("%"):
            return default  # for height, % doesn't map cleanly
        try:
            return float(s)
        except Exception:
            return default

    def _substitute_vars(self, text: str) -> str:
        brand = self.product["brands"]["wevac"]
        warranty = self.product.get("warranty", {})
        product = self.product["product"]
        repl = {
            "{{brand.display_name}}": brand["display_name_cn"],
            "{{brand.support_email}}": brand["support_email"],
            "{{brand.website}}": brand["website"],
            "{{warranty.years}}": str(warranty.get("years", 2)),
            "{{localized.product_name}}": product["name_cn"],
        }
        for k, v in repl.items():
            text = text.replace(k, v)
        return text


def main():
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else (HERE / "iter-01" / "output.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    builder = ManualBuilder(out)
    builder.build()


if __name__ == "__main__":
    main()
