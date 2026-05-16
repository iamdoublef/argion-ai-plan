from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image


ROOT = Path(r"D:\work\private\yjsplan\research\yjs-manual-opt\swiss")
HTML_PATH = ROOT / "output" / "imt050-wevac-eu-cn.html"
IMAGE_ROOT = ROOT / "output"

PAGE_W_MM = 148
PAGE_H_MM = 210
MARGIN_MM = 10
CONTENT_W_MM = PAGE_W_MM - MARGIN_MM * 2

RED = "E63946"
BLACK = "000000"
DARK = "1A1A1A"
GRAY = "8E8E93"
LIGHT_GRAY = "F2F2F7"
ZEBRA_GRAY = "F2F2F7"
BORDER_GRAY = "CCCCCC"
INFO_BLUE = "007AFF"

LATIN_FONT = "Arial"
CJK_FONT = "Microsoft YaHei"
LATIN_BOLD_FONT = "Arial Black"
MONO_FONT = "Courier New"
CHAR_SPACING_TWIPS = 5
WARNING_CHAR_SPACING_TWIPS = 8

BODY_PT = 7.05
SECTION_TITLE_PT = 9.0
SUB_TITLE_PT = 7.5
CHAPTER_NUM_PT = 13.5
CHAPTER_TITLE_PT = 11.0
TABLE_BODY_PT = 6.7
TABLE_HEADER_PT = 6.0
SMALL_PT = 5.4

BODY_LINE_SPACING = 1.16
BULLET_LINE_SPACING = 1.13
BULLET_SPACE_AFTER_PT = 1.6
STEP_LINE_SPACING = 1.10
TABLE_LINE_SPACING = 1.00
TABLE_CELL_PAD_NORMAL = 56
TABLE_CELL_PAD_COMPACT = 32
ALERT_CELL_PAD_V = 46


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER_GRAY, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_horizontal_borders(cell, color: str = BORDER_GRAY, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "bottom"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
    for edge in ("left", "right", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_bottom_border(cell, color: str = RED, size: int = 10) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "right"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)


def set_cell_top_border(cell, color: str = "EEEEEE", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("left", "right", "bottom"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "nil")
    top = borders.find(qn("w:top"))
    if top is None:
        top = OxmlElement("w:top")
        borders.append(top)
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color)


def set_cell_margins(cell, top=50, start=70, bottom=50, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_height(row, height_twips: int, rule: str = "atLeast") -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), rule)


def set_cell_width(cell, width_mm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_mm * 56.7)))


def set_paragraph_border(paragraph, *, left=None, bottom=None, top=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name, spec in (("left", left), ("bottom", bottom), ("top", top)):
        if not spec:
            continue
        edge = p_bdr.find(qn("w:" + edge_name))
        if edge is None:
            edge = OxmlElement("w:" + edge_name)
            p_bdr.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(spec.get("size", 8)))
        edge.set(qn("w:space"), str(spec.get("space", 2)))
        edge.set(qn("w:color"), spec.get("color", BLACK))


def set_table_width(table, width_mm: float = CONTENT_W_MM) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths_mm: list[float]) -> None:
    if not widths_mm:
        return
    tbl = table._tbl
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_mm in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_mm * 56.7)))
        grid.append(col)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_run_fonts(run, latin=LATIN_FONT, east_asia=CJK_FONT, latin_bold=LATIN_BOLD_FONT, is_bold=False, is_mono=False):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    if is_mono:
        rFonts.set(qn("w:ascii"), MONO_FONT)
        rFonts.set(qn("w:hAnsi"), MONO_FONT)
        rFonts.set(qn("w:cs"), MONO_FONT)
    else:
        rFonts.set(qn("w:ascii"), latin_bold if is_bold else latin)
        rFonts.set(qn("w:hAnsi"), latin_bold if is_bold else latin)
        rFonts.set(qn("w:cs"), latin_bold if is_bold else latin)
    rFonts.set(qn("w:eastAsia"), east_asia)


def set_run_font(run, size=BODY_PT, bold=False, color=BLACK, is_mono=False) -> None:
    run.font.name = MONO_FONT if is_mono else (LATIN_BOLD_FONT if bold else LATIN_FONT)
    set_run_fonts(run, is_bold=bold, is_mono=is_mono)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None
    run._element.get_or_add_rPr()
    color_node = run._element.rPr.find(qn("w:color"))
    if color_node is None:
        color_node = OxmlElement("w:color")
        run._element.rPr.append(color_node)
    color_node.set(qn("w:val"), color)
    if not is_mono:
        spacing_node = run._element.rPr.find(qn("w:spacing"))
        if spacing_node is None:
            spacing_node = OxmlElement("w:spacing")
            run._element.rPr.append(spacing_node)
        spacing_node.set(qn("w:val"), str(CHAR_SPACING_TWIPS))


def override_run_spacing(run, twips: int) -> None:
    r_pr = run._element.get_or_add_rPr()
    spacing_node = r_pr.find(qn("w:spacing"))
    if spacing_node is None:
        spacing_node = OxmlElement("w:spacing")
        r_pr.append(spacing_node)
    spacing_node.set(qn("w:val"), str(twips))


def set_run_position(run, half_points: int) -> None:
    r_pr = run._element.get_or_add_rPr()
    pos = r_pr.find(qn("w:position"))
    if pos is None:
        pos = OxmlElement("w:position")
        r_pr.append(pos)
    pos.set(qn("w:val"), str(half_points))


def set_run_shading(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def compact_paragraph(paragraph, size=BODY_PT, after=2, before=0, line=BODY_LINE_SPACING) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=size)


def add_text_runs(paragraph, node, size=BODY_PT, color=BLACK, bold=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node).replace("\n", " ")
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)
        return
    if not isinstance(node, Tag):
        return
    next_bold = bold or node.name in ("b", "strong")
    for child in node.children:
        add_text_runs(paragraph, child, size=size, color=color, bold=next_bold)


def add_para(doc_or_cell, text_or_node, *, size=BODY_PT, bold=False, color=BLACK, align=None, after=2, before=0, is_mono=False):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if isinstance(text_or_node, Tag):
        add_text_runs(p, text_or_node, size=size, color=color, bold=bold)
    else:
        run = p.add_run(str(text_or_node))
        set_run_font(run, size=size, bold=bold, color=color, is_mono=is_mono)
    compact_paragraph(p, size=size, after=after, before=before)
    return p


def parse_mm(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)mm".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def parse_px(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)px".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def parse_pct(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)%".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def table_column_widths(rows: list[Tag], max_cols: int) -> list[float]:
    if not rows or max_cols <= 0:
        return []
    first_cells = rows[0].find_all(["td", "th"], recursive=False)
    widths: list[float | None] = []
    for idx in range(max_cols):
        style = first_cells[idx].get("style", "") if idx < len(first_cells) else ""
        widths.append(parse_pct(style, "width"))
    used = sum(w for w in widths if w is not None)
    missing = sum(1 for w in widths if w is None)
    fallback = max(0.0, 100.0 - used) / missing if missing else 0.0
    pct_widths = [(w if w is not None else fallback) for w in widths]
    total = sum(pct_widths) or 100.0
    return [CONTENT_W_MM * w / total for w in pct_widths]


def image_path(img: Tag) -> Path:
    src = img.get("src", "").replace("./", "")
    return IMAGE_ROOT / src


def fit_image_mm(path: Path, max_w_mm: float, max_h_mm: float) -> tuple[float, float]:
    with Image.open(path) as im:
        w, h = im.size
    scale = min(max_w_mm / w, max_h_mm / h)
    return max(1, w * scale), max(1, h * scale)


def add_image_paragraph(parent, img: Tag, *, align=WD_ALIGN_PARAGRAPH.CENTER, force_w=None, force_h=None):
    path = image_path(img)
    if not path.exists():
        return None
    style = img.get("style", "")
    alt = img.get("alt", "")
    is_warranty_separator = "Warranty separator" in alt
    if is_warranty_separator:
        return None
    max_h = force_h or parse_mm(style, "max-height") or (4.0 if is_warranty_separator else 45)
    pct_w = parse_pct(style, "max-width")
    max_w = force_w or (60.0 if is_warranty_separator else (CONTENT_W_MM * pct_w / 100 if pct_w else CONTENT_W_MM))
    px_h = parse_px(style, "height")
    if px_h:
        max_h = px_h * 0.2645
    w_mm, h_mm = fit_image_mm(path, max_w, max_h)
    p = parent.add_paragraph()
    p.alignment = align
    if is_warranty_separator:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    else:
        p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))
    return p


def add_header_strip(doc: Document, node: Tag, *, after_pt=4) -> None:
    p = doc.add_paragraph()
    set_paragraph_border(p, top={"size": 18, "color": BLACK, "space": 1})
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = 1
    text = node.get_text(" ", strip=True)
    left = text.split("CH.")[0].strip() if "CH." in text else " ".join(text.split()[:1])
    right = "CH." + text.split("CH.", 1)[1].strip() if "CH." in text else text.replace(left, "").strip()
    r1 = p.add_run(left)
    set_run_font(r1, size=6.75, bold=True)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(CONTENT_W_MM), WD_ALIGN_PARAGRAPH.RIGHT)
    r_tab = p.add_run("\t")
    set_run_font(r_tab, size=SMALL_PT, color=GRAY, is_mono=True)
    r2 = p.add_run(right)
    set_run_font(r2, size=SMALL_PT, color=GRAY, is_mono=True)


def add_section_title(doc: Document, node: Tag) -> None:
    p = doc.add_paragraph()
    set_paragraph_border(p, left={"size": 18, "color": BLACK, "space": 4})
    p.paragraph_format.left_indent = Mm(1.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1
    num = node.select_one(".chapter-num")
    num_text = num.get_text(strip=True) if num else ""
    full_text = node.get_text(" ", strip=True)
    title = full_text.replace(num_text, "", 1).strip()
    r1 = p.add_run(num_text + " ")
    set_run_font(r1, size=CHAPTER_NUM_PT, bold=True, color=RED)
    r2 = p.add_run(title)
    set_run_font(r2, size=CHAPTER_TITLE_PT, bold=True, color=BLACK)


def add_sub_title(doc: Document, node: Tag, *, before=4) -> None:
    p = add_para(doc, node.get_text(" ", strip=True), size=SUB_TITLE_PT, bold=True, after=4, before=before)
    set_paragraph_border(p, bottom={"size": 6, "color": BLACK, "space": 1})


def add_bullet_list(
    parent,
    ul: Tag,
    *,
    size=BODY_PT,
    red_bullet=True,
    tight=False,
    indent_mm=3.2,
    line_spacing=None,
    space_after_pt=None,
    char_spacing_twips=None,
) -> None:
    for li in ul.find_all("li", recursive=False):
        p = parent.add_paragraph()
        p.paragraph_format.left_indent = Mm(indent_mm)
        p.paragraph_format.first_line_indent = Mm(-indent_mm)
        p.paragraph_format.space_before = Pt(0)
        after_pt = space_after_pt if space_after_pt is not None else (1.35 if tight else BULLET_SPACE_AFTER_PT)
        p.paragraph_format.space_after = Pt(after_pt)
        p.paragraph_format.line_spacing = line_spacing if line_spacing is not None else (0.96 if tight else BULLET_LINE_SPACING)
        bullet = p.add_run("•    ")
        bullet_size = 5.25 if tight else 5.8
        set_run_font(bullet, size=bullet_size, bold=True, color=RED if red_bullet else BLACK)
        add_text_runs(p, li, size=size, color=BLACK)
        if char_spacing_twips is not None:
            for run in p.runs:
                override_run_spacing(run, char_spacing_twips)


def add_alert_box(doc: Document, node: Tag, *, compact_safety=False) -> None:
    classes = set(node.get("class", []))
    is_note = "note-box" in classes
    title = node.select_one(".box-title")
    if is_note and title and "WEEE" in title.get_text(" ", strip=True):
        spacer = doc.add_paragraph()
        compact_paragraph(spacer, size=BODY_PT, after=3, line=0.1)
    color = RED if "warning-box" in classes else (BLACK if "caution-box" in classes else GRAY)
    fill = "FFFFFF" if not is_note else LIGHT_GRAY
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if not is_note:
        set_cell_borders(cell, color=color, size=12 if color == RED else 8)
    pad_start = 205 if is_note else 176
    pad_v = 44 if is_note else ALERT_CELL_PAD_V
    set_cell_margins(cell, top=pad_v, start=pad_start, bottom=pad_v, end=110)
    if fill != "FFFFFF":
        set_cell_shading(cell, fill)
    if title:
        p_title = cell.paragraphs[0]
        r_title = p_title.add_run(title.get_text(" ", strip=True))
        title_size = 6.38 if compact_safety else 6.5
        set_run_font(r_title, size=title_size, bold=True, color=color if color != GRAY else BLACK)
        compact_paragraph(p_title, size=title_size, after=1 if compact_safety else 2)
    for img in node.find_all("img"):
        icon_h = 5.2 if compact_safety else 6
        icon_w = 6.9 if compact_safety else 8
        p_img = add_image_paragraph(cell, img, align=WD_ALIGN_PARAGRAPH.LEFT, force_h=icon_h, force_w=icon_w)
        if p_img:
            if compact_safety:
                p_img.paragraph_format.left_indent = Mm(-3.1)
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(1)
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                add_para(cell, text, size=BODY_PT, after=1)
            continue
        if not isinstance(child, Tag):
            continue
        if child == title or child.name == "img":
            continue
        if child.name == "ul":
            add_bullet_list(
                cell,
                child,
                size=6.98 if compact_safety else BODY_PT,
                tight=compact_safety or is_note,
                indent_mm=4.7,
                line_spacing=1.10 if "warning-box" in classes and not compact_safety else None,
                space_after_pt=0.5 if "warning-box" in classes and not compact_safety else None,
                char_spacing_twips=WARNING_CHAR_SPACING_TWIPS if "warning-box" in classes else None,
            )
        elif child.name == "p":
            add_para(cell, child, size=BODY_PT, after=1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image_row(parent, node: Tag, *, max_h_override=None) -> None:
    children = [c for c in node.children if isinstance(c, Tag) and (c.name == "img" or c.get_text(strip=True))]
    imgs = node.find_all("img", recursive=False)
    if not imgs and node.name != "div":
        return
    if "step-figures" in node.get("class", []) and len(imgs) == 2:
        p = parent.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Mm(30)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1
        for i, img in enumerate(imgs):
            path = image_path(img)
            if not path.exists():
                continue
            style = img.get("style", "")
            max_h = max_h_override or parse_mm(style, "max-height") or 24
            pct_w = parse_pct(style, "max-width")
            max_w = CONTENT_W_MM * pct_w / 100 if pct_w else 45
            w_mm, h_mm = fit_image_mm(path, max_w, max_h)
            run = p.add_run()
            run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))
            if i == 1:
                set_run_position(run, 41)
            if i != len(imgs) - 1:
                gap = p.add_run("       ")
                set_run_font(gap, size=7)
        return
    if "status-indicator-row" in node.get("class", []):
        p = parent.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        for child in children:
            if child.name == "span":
                r = p.add_run(child.get_text(" ", strip=True) + " ")
                set_run_font(r, size=6.15, bold=False)
            elif child.name == "img":
                path = image_path(child)
                if path.exists():
                    run = p.add_run()
                    run.add_picture(str(path), height=Mm(4))
                p.add_run(" ")
        return
    cols = max(1, min(len(imgs), 3))
    table = parent.add_table(rows=1, cols=cols)
    set_table_width(table, CONTENT_W_MM)
    remove_table_borders(table)
    for i, img in enumerate(imgs):
        cell = table.cell(0, i)
        set_cell_margins(cell, top=0, start=20, bottom=0, end=20)
        path = image_path(img)
        if not path.exists():
            continue
        style = img.get("style", "")
        max_h = max_h_override or parse_mm(style, "max-height") or 26
        pct_w = parse_pct(style, "max-width")
        max_w = CONTENT_W_MM / cols - 2
        if pct_w:
            max_w = min(max_w, CONTENT_W_MM * pct_w / 100)
        w_mm, h_mm = fit_image_mm(path, max_w, max_h)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))


def add_html_table(doc: Document, node: Tag, *, compact_warranty=False, troubleshooting=False) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    max_cols = max(len(r.find_all(["td", "th"], recursive=False)) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = table_column_widths(rows, max_cols)
    set_table_grid(table, col_widths)
    is_warranty_card = "warranty-card" in node.get("class", [])
    compact = len(rows) > 8 or is_warranty_card
    size = 6.52 if compact and max_cols >= 3 else TABLE_BODY_PT
    remove_table_borders(table)
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"], recursive=False)
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            if c_idx < len(col_widths):
                set_cell_width(cell, col_widths[c_idx])
            if compact_warranty or troubleshooting:
                set_cell_borders(cell, color="D9D9D9", size=4)
            else:
                set_cell_horizontal_borders(cell)
            pad_v = 32 if troubleshooting and compact else (52 if compact_warranty and not compact else (TABLE_CELL_PAD_COMPACT if compact else 36))
            start_pad = 87 if compact_warranty else 55
            set_cell_margins(cell, top=pad_v, start=start_pad, bottom=pad_v, end=55)
            if is_warranty_card:
                set_cell_shading(cell, ZEBRA_GRAY if r_idx % 2 == 1 else "FFFFFF")
            elif r_idx == 0:
                set_cell_shading(cell, DARK)
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "FFFFFF")
            else:
                set_cell_shading(cell, ZEBRA_GRAY)
            text = cells[c_idx].get_text("\n", strip=True) if c_idx < len(cells) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = TABLE_LINE_SPACING
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run_size = TABLE_HEADER_PT if r_idx == 0 else size
            set_run_font(
                run,
                size=run_size,
                bold=(r_idx == 0 and not is_warranty_card),
                color="FFFFFF" if r_idx == 0 and not is_warranty_card else DARK,
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP if compact_warranty else WD_ALIGN_VERTICAL.CENTER
        row_height = 242 if compact_warranty and not compact else (225 if compact else 215)
        set_row_height(table.rows[r_idx], row_height)

def add_step_flow(doc: Document, node: Tag) -> None:
    for step in node.select(".step-flow-step"):
        row = step.select_one(".step-flow-row")
        if row:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = STEP_LINE_SPACING
            num = row.select_one(".step-num").get_text(strip=True)
            r_num = p.add_run(" " + num + " ")
            set_run_font(r_num, size=BODY_PT, bold=True, color="FFFFFF")
            set_run_shading(r_num, BLACK)
            r_gap = p.add_run("   ")
            set_run_font(r_gap, size=BODY_PT)
            add_text_runs(p, row.select_one(".step-text"), size=BODY_PT)
        for fig in step.select(".figure-row"):
            add_image_row(doc, fig)


def add_cover(doc: Document, page: Tag) -> None:
    brand = page.select_one(".cover-brand").get_text(" ", strip=True)
    p_brand = doc.add_paragraph()
    p_brand.paragraph_format.space_before = Pt(20)
    p_brand.paragraph_format.space_after = Pt(138)
    # PDF target uses a SHORT red dash before \u5a01\u5bcc\u53ef, not a long line; one character matches better
    r_line = p_brand.add_run(u"\u2501\u2501 ")
    set_run_font(r_line, size=7.5, bold=True, color=RED)
    r_brand = p_brand.add_run(brand)
    set_run_font(r_brand, size=7.5, bold=True, color=BLACK)
    img = page.find("img")
    if img:
        p_img = add_image_paragraph(doc, img, align=WD_ALIGN_PARAGRAPH.LEFT, force_h=32.5, force_w=42)
        if p_img:
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(30)
    model = page.select_one(".cover-model").get_text(" ", strip=True)
    title = page.select_one(".cover-title").get_text(" ", strip=True)
    subtitle = page.select_one(".cover-subtitle").get_text(" ", strip=True)
    add_para(doc, model, size=6.0, bold=True, color=RED, after=2, is_mono=True)
    add_para(doc, title, size=18.0, bold=True, color=DARK, after=0)
    add_para(doc, subtitle, size=7.5, color=GRAY, after=1)
    p_div = doc.add_paragraph()
    r_div = p_div.add_run(u"\u2501\u2501\u2501\u2501")
    set_run_font(r_div, size=7.0, bold=True, color=RED)
    p_div.paragraph_format.space_after = Pt(113)
    bottom = page.select_one(".cover-bottom").get_text(" ", strip=True)
    p_rule = doc.add_paragraph()
    set_paragraph_border(p_rule, top={"size": 10, "color": BLACK, "space": 1})
    p_rule.paragraph_format.space_after = Pt(3)
    add_para(doc, bottom, size=SMALL_PT, color=GRAY, after=0)


def add_toc_page(doc: Document, page: Tag) -> None:
    add_header_strip(doc, page.select_one(".header-strip"))
    add_para(doc, page.select_one(".toc-title").get_text(" ", strip=True), size=15, bold=True, after=10)
    for item in page.select(".toc-item"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.tab_stops.add_tab_stop(Mm(CONTENT_W_MM), WD_ALIGN_PARAGRAPH.RIGHT)
        chap = item.select_one(".toc-chapter").get_text(strip=True)
        name = item.select_one(".toc-name").get_text(strip=True)
        page_no = item.select_one(".toc-page").get_text(strip=True)
        r1 = p.add_run(chap + "  ")
        set_run_font(r1, size=6.75, bold=True, color=RED, is_mono=True)
        r2 = p.add_run(name)
        set_run_font(r2, size=7.5, bold=True)
        r3 = p.add_run("\t" + page_no)
        set_run_font(r3, size=6.38, color=GRAY, is_mono=True)


def add_body_page(doc: Document, page: Tag) -> None:
    page_classes = set(page.get("class", []))
    previous_was_list = False
    for child in page.find_all(recursive=False):
        if not isinstance(child, Tag):
            continue
        after_list = previous_was_list
        previous_was_list = child.name == "ul"
        classes = set(child.get("class", []))
        if "page-footer" in classes:
            continue
        if "header-strip" in classes:
            add_header_strip(doc, child, after_pt=8.8)
        elif "section-title" in classes:
            add_section_title(doc, child)
        elif "sub-title" in classes:
            add_sub_title(doc, child, before=8 if after_list else 4)
        elif child.name == "p":
            text = child.get_text(" ", strip=True)
            if "compact-warranty" in page_classes and "support@wevactech.com |" in text:
                text = text.replace("联系我们： ", "联系我们：\n")
            add_para(doc, text if text else child, size=BODY_PT, color=DARK, after=3)
        elif child.name == "ul":
            if "compact-warranty" in page_classes:
                add_bullet_list(doc, child, line_spacing=1.0, space_after_pt=0.25)
            else:
                add_bullet_list(doc, child)
        elif "warning-box" in classes or "caution-box" in classes or "note-box" in classes:
            if "compact-ts" in page_classes and "caution-box" in classes:
                spacer = doc.add_paragraph()
                compact_paragraph(spacer, size=BODY_PT, after=7, line=0.1)
            add_alert_box(doc, child, compact_safety=("compact-safety" in page_classes))
        elif "step-flow" in classes:
            add_step_flow(doc, child)
        elif "fig-wrap" in classes:
            img = child.find("img")
            if img:
                add_image_paragraph(doc, img)
        elif "figure-row" in classes:
            add_image_row(doc, child)
        elif child.name == "table":
            add_html_table(
                doc,
                child,
                compact_warranty=("compact-warranty" in page_classes),
                troubleshooting=("troubleshooting-primary" in page_classes),
            )


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=SMALL_PT, color=GRAY, is_mono=True)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def configure_section(section, page_no: int | None = None) -> None:
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(10.2)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(3.5)
    section.different_first_page_header_footer = False
    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = ""
    if page_no is None:
        return
    table = footer.add_table(rows=1, cols=2, width=Mm(CONTENT_W_MM))
    set_table_width(table)
    remove_table_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_width(left, 92)
    set_cell_width(right, CONTENT_W_MM - 92)
    for cell in (left, right):
        set_cell_top_border(cell)
        set_cell_margins(cell, top=35, start=0, bottom=0, end=0)
    p_left = left.paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    r = p_left.add_run("威富可 IMT050 说明书")
    set_run_font(r, size=SMALL_PT, color=GRAY)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    r_no = p_right.add_run(str(page_no))
    set_run_font(r_no, size=SMALL_PT, color=GRAY, is_mono=True)


def setup_document() -> Document:
    doc = Document()
    configure_section(doc.sections[0])
    return doc


def build(out_path: Path) -> None:
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    pages = soup.select(".page")
    doc = setup_document()
    for idx, page in enumerate(pages):
        if idx == 0:
            add_cover(doc, page)
        elif idx == 1:
            add_toc_page(doc, page)
        else:
            add_body_page(doc, page)
        if idx != len(pages) - 1:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, page_no=idx + 2)
    doc.save(out_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("out", type=Path)
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)
    build(args.out)
