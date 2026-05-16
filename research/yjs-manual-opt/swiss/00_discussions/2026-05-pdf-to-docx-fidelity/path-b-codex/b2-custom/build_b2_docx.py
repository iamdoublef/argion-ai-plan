from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image


ROOT = Path(r"D:\work\private\yjsplan\research\yjs-manual-opt\swiss")
HTML_PATH = ROOT / "output" / "imt050-wevac-eu-cn.html"
IMAGE_ROOT = ROOT / "output"

PAGE_W_MM = 148
PAGE_H_MM = 210
MARGIN_MM = 10
CONTENT_W_MM = PAGE_W_MM - MARGIN_MM * 2

RED = "E63946"
BLACK = "000000"
GRAY = "8E8E93"
LIGHT_GRAY = "F2F2F7"
BORDER_GRAY = "CCCCCC"

LATIN_FONT = "Arial"
CJK_FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER_GRAY, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=50, start=70, bottom=50, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, *, left=None, bottom=None, top=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name, spec in (("left", left), ("bottom", bottom), ("top", top)):
        if not spec:
            continue
        edge = p_bdr.find(qn("w:" + edge_name))
        if edge is None:
            edge = OxmlElement("w:" + edge_name)
            p_bdr.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(spec.get("size", 8)))
        edge.set(qn("w:space"), str(spec.get("space", 2)))
        edge.set(qn("w:color"), spec.get("color", BLACK))


def set_table_width(table, width_mm: float = CONTENT_W_MM) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_run_font(run, size=9, bold=False, color=BLACK) -> None:
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = None
    run._element.get_or_add_rPr()
    color_node = run._element.rPr.find(qn("w:color"))
    if color_node is None:
        color_node = OxmlElement("w:color")
        run._element.rPr.append(color_node)
    color_node.set(qn("w:val"), color)


def set_run_shading(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def compact_paragraph(paragraph, size=9, after=2, before=0, line=1.08) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        if run.font.size is None:
            set_run_font(run, size=size)


def add_text_runs(paragraph, node, size=9, color=BLACK, bold=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node).replace("\n", " ")
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)
        return
    if not isinstance(node, Tag):
        return
    next_bold = bold or node.name in ("b", "strong")
    for child in node.children:
        add_text_runs(paragraph, child, size=size, color=color, bold=next_bold)


def add_para(doc_or_cell, text_or_node, *, size=9, bold=False, color=BLACK, align=None, after=2, before=0):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if isinstance(text_or_node, Tag):
        add_text_runs(p, text_or_node, size=size, color=color, bold=bold)
    else:
        run = p.add_run(str(text_or_node))
        set_run_font(run, size=size, bold=bold, color=color)
    compact_paragraph(p, size=size, after=after, before=before)
    return p


def parse_mm(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)mm".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def parse_px(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)px".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def parse_pct(style: str, key: str, default: float | None = None) -> float | None:
    m = re.search(r"{}:\s*([0-9.]+)%".format(re.escape(key)), style)
    return float(m.group(1)) if m else default


def image_path(img: Tag) -> Path:
    src = img.get("src", "").replace("./", "")
    return IMAGE_ROOT / src


def fit_image_mm(path: Path, max_w_mm: float, max_h_mm: float) -> tuple[float, float]:
    with Image.open(path) as im:
        w, h = im.size
    scale = min(max_w_mm / w, max_h_mm / h)
    return max(1, w * scale), max(1, h * scale)


def add_image_paragraph(parent, img: Tag, *, align=WD_ALIGN_PARAGRAPH.CENTER, force_w=None, force_h=None):
    path = image_path(img)
    if not path.exists():
        return None
    style = img.get("style", "")
    max_h = force_h or parse_mm(style, "max-height") or 45
    pct_w = parse_pct(style, "max-width")
    max_w = force_w or (CONTENT_W_MM * pct_w / 100 if pct_w else CONTENT_W_MM)
    px_h = parse_px(style, "height")
    if px_h:
        max_h = px_h * 0.2645
    w_mm, h_mm = fit_image_mm(path, max_w, max_h)
    p = parent.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))
    return p


def add_header_strip(doc: Document, node: Tag) -> None:
    p = doc.add_paragraph()
    set_paragraph_border(p, top={"size": 18, "color": BLACK, "space": 1})
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1
    text = node.get_text(" ", strip=True)
    left = text.split("CH.")[0].strip() if "CH." in text else " ".join(text.split()[:1])
    right = "CH." + text.split("CH.", 1)[1].strip() if "CH." in text else text.replace(left, "").strip()
    r1 = p.add_run(left)
    set_run_font(r1, size=8.5, bold=True)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(CONTENT_W_MM), WD_ALIGN_PARAGRAPH.RIGHT)
    r_tab = p.add_run("\t")
    set_run_font(r_tab, size=7.2, color=GRAY)
    r2 = p.add_run(right)
    set_run_font(r2, size=7.2, color=GRAY)


def add_section_title(doc: Document, node: Tag) -> None:
    p = doc.add_paragraph()
    set_paragraph_border(p, left={"size": 18, "color": BLACK, "space": 4})
    p.paragraph_format.left_indent = Mm(1.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1
    num = node.select_one(".chapter-num")
    num_text = num.get_text(strip=True) if num else ""
    full_text = node.get_text(" ", strip=True)
    title = full_text.replace(num_text, "", 1).strip()
    r1 = p.add_run(num_text + " ")
    set_run_font(r1, size=18, bold=True, color=RED)
    r2 = p.add_run(title)
    set_run_font(r2, size=16, bold=True, color=BLACK)


def add_sub_title(doc: Document, node: Tag) -> None:
    p = add_para(doc, node.get_text(" ", strip=True), size=8.5, bold=True, after=4, before=4)
    set_paragraph_border(p, bottom={"size": 6, "color": BLACK, "space": 1})


def add_bullet_list(parent, ul: Tag, *, size=8.6, red_bullet=True) -> None:
    for li in ul.find_all("li", recursive=False):
        p = parent.add_paragraph()
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.first_line_indent = Mm(-3)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = 1.05
        bullet = p.add_run(u"\u2022  ")
        set_run_font(bullet, size=size, bold=True, color=RED if red_bullet else BLACK)
        add_text_runs(p, li, size=size, color=BLACK)


def add_alert_box(doc: Document, node: Tag) -> None:
    classes = set(node.get("class", []))
    color = RED if "warning-box" in classes else (BLACK if "caution-box" in classes else GRAY)
    fill = "FFFFFF" if "note-box" not in classes else LIGHT_GRAY
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_borders(cell, color=color, size=12 if color == RED else 8)
    set_cell_margins(cell, top=70, start=110, bottom=70, end=110)
    if fill != "FFFFFF":
        set_cell_shading(cell, fill)
    title = node.select_one(".box-title")
    if title:
        add_para(cell, title.get_text(" ", strip=True), size=8.3, bold=True, color=color if color != GRAY else BLACK, after=2)
    for img in node.find_all("img"):
        add_image_paragraph(cell, img, align=WD_ALIGN_PARAGRAPH.LEFT, force_h=6, force_w=8)
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                add_para(cell, text, size=8.1, after=1)
            continue
        if not isinstance(child, Tag):
            continue
        if child == title or child.name == "img":
            continue
        if child.name == "ul":
            add_bullet_list(cell, child, size=7.85)
        elif child.name == "p":
            add_para(cell, child, size=8.1, after=1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image_row(parent, node: Tag, *, max_h_override=None) -> None:
    children = [c for c in node.children if isinstance(c, Tag) and (c.name == "img" or c.get_text(strip=True))]
    imgs = node.find_all("img", recursive=False)
    if not imgs and node.name != "div":
        return
    if "status-indicator-row" in node.get("class", []):
        p = parent.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        for child in children:
            if child.name == "span":
                r = p.add_run(child.get_text(" ", strip=True) + " ")
                set_run_font(r, size=8, bold=True)
            elif child.name == "img":
                path = image_path(child)
                if path.exists():
                    run = p.add_run()
                    run.add_picture(str(path), height=Mm(4))
                p.add_run(" ")
        return
    cols = max(1, min(len(imgs), 3))
    table = parent.add_table(rows=1, cols=cols)
    set_table_width(table, CONTENT_W_MM)
    remove_table_borders(table)
    for i, img in enumerate(imgs):
        cell = table.cell(0, i)
        set_cell_margins(cell, top=0, start=20, bottom=0, end=20)
        path = image_path(img)
        if not path.exists():
            continue
        style = img.get("style", "")
        max_h = max_h_override or parse_mm(style, "max-height") or 26
        pct_w = parse_pct(style, "max-width")
        max_w = CONTENT_W_MM / cols - 2
        if pct_w:
            max_w = min(max_w, CONTENT_W_MM * pct_w / 100)
        w_mm, h_mm = fit_image_mm(path, max_w, max_h)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))
    parent.add_paragraph().paragraph_format.space_after = Pt(1)


def add_html_table(doc: Document, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    max_cols = max(len(r.find_all(["td", "th"], recursive=False)) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    compact = len(rows) > 10 or "warranty-card" in node.get("class", [])
    size = 6.9 if compact and max_cols >= 3 else (7.25 if compact else 7.8)
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"], recursive=False)
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            set_cell_borders(cell)
            set_cell_margins(cell, top=35 if compact else 50, start=55, bottom=35 if compact else 50, end=55)
            if r_idx == 0:
                set_cell_shading(cell, BLACK)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            text = cells[c_idx].get_text("\n", strip=True) if c_idx < len(cells) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 and max_cols <= 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=size, bold=(r_idx == 0), color="FFFFFF" if r_idx == 0 else BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_step_flow(doc: Document, node: Tag) -> None:
    for step in node.select(".step-flow-step"):
        row = step.select_one(".step-flow-row")
        if row:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.04
            num = row.select_one(".step-num").get_text(strip=True)
            r_num = p.add_run("  " + num + "  ")
            set_run_font(r_num, size=8.4, bold=True, color="FFFFFF")
            set_run_shading(r_num, BLACK)
            r_gap = p.add_run("  ")
            set_run_font(r_gap, size=8.4)
            add_text_runs(p, row.select_one(".step-text"), size=8.4)
        for fig in step.select(".figure-row"):
            add_image_row(doc, fig)


def add_cover(doc: Document, page: Tag) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(75)
    r_line = p.add_run(u"\u2501\u2501\u2501 ")
    set_run_font(r_line, size=9, bold=True, color=RED)
    brand = page.select_one(".cover-brand").get_text(" ", strip=True)
    r_brand = p.add_run(brand)
    set_run_font(r_brand, size=10.5, bold=True)
    img = page.find("img")
    if img:
        p_img = add_image_paragraph(doc, img, align=WD_ALIGN_PARAGRAPH.LEFT, force_h=34, force_w=42)
        if p_img:
            p_img.paragraph_format.space_after = Pt(30)
    model = page.select_one(".cover-model").get_text(" ", strip=True)
    title = page.select_one(".cover-title").get_text(" ", strip=True)
    subtitle = page.select_one(".cover-subtitle").get_text(" ", strip=True)
    add_para(doc, model, size=7.4, bold=True, color=RED, after=2)
    add_para(doc, title, size=22, bold=True, after=0)
    add_para(doc, subtitle, size=10, color=GRAY, after=1)
    p_div = doc.add_paragraph()
    r_div = p_div.add_run(u"\u2501\u2501\u2501\u2501")
    set_run_font(r_div, size=9, bold=True, color=RED)
    p_div.paragraph_format.space_after = Pt(92)
    bottom = page.select_one(".cover-bottom").get_text(" ", strip=True)
    p_rule = doc.add_paragraph()
    set_paragraph_border(p_rule, top={"size": 10, "color": BLACK, "space": 1})
    p_rule.paragraph_format.space_after = Pt(3)
    add_para(doc, bottom, size=6.8, color=GRAY, after=0)


def add_toc_page(doc: Document, page: Tag) -> None:
    add_header_strip(doc, page.select_one(".header-strip"))
    add_para(doc, page.select_one(".toc-title").get_text(" ", strip=True), size=18, bold=True, after=10)
    for item in page.select(".toc-item"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.tab_stops.add_tab_stop(Mm(CONTENT_W_MM), WD_ALIGN_PARAGRAPH.RIGHT)
        chap = item.select_one(".toc-chapter").get_text(strip=True)
        name = item.select_one(".toc-name").get_text(strip=True)
        page_no = item.select_one(".toc-page").get_text(strip=True)
        r1 = p.add_run(chap + "  ")
        set_run_font(r1, size=11, bold=True, color=RED)
        r2 = p.add_run(name)
        set_run_font(r2, size=10, bold=True)
        r3 = p.add_run("\t" + page_no)
        set_run_font(r3, size=10, color=GRAY)


def add_body_page(doc: Document, page: Tag) -> None:
    for child in page.find_all(recursive=False):
        if not isinstance(child, Tag):
            continue
        classes = set(child.get("class", []))
        if "page-footer" in classes:
            continue
        if "header-strip" in classes:
            add_header_strip(doc, child)
        elif "section-title" in classes:
            add_section_title(doc, child)
        elif "sub-title" in classes:
            add_sub_title(doc, child)
        elif child.name == "p":
            add_para(doc, child, size=9, after=3)
        elif child.name == "ul":
            add_bullet_list(doc, child)
        elif "warning-box" in classes or "caution-box" in classes or "note-box" in classes:
            add_alert_box(doc, child)
        elif "step-flow" in classes:
            add_step_flow(doc, child)
        elif "fig-wrap" in classes:
            img = child.find("img")
            if img:
                add_image_paragraph(doc, img)
        elif "figure-row" in classes:
            add_image_row(doc, child)
        elif child.name == "table":
            add_html_table(doc, child)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    section.top_margin = Mm(MARGIN_MM)
    section.bottom_margin = Mm(MARGIN_MM)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(5)
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Mm(CONTENT_W_MM), WD_ALIGN_PARAGRAPH.RIGHT)
    set_paragraph_border(p, top={"size": 4, "color": "EEEEEE", "space": 1})
    r = p.add_run("威富可 IMT050 说明书")
    set_run_font(r, size=6.2, color=GRAY)
    r_tab = p.add_run("\t")
    set_run_font(r_tab, size=6.2, color=GRAY)
    add_page_field(p)
    for run in p.runs:
        if run.font.size is None:
            set_run_font(run, size=6.2, color=GRAY)
    return doc


def build(out_path: Path) -> None:
    soup = BeautifulSoup(HTML_PATH.read_text(encoding="utf-8"), "html.parser")
    pages = soup.select(".page")
    doc = setup_document()
    for idx, page in enumerate(pages):
        if idx == 0:
            add_cover(doc, page)
        elif idx == 1:
            add_toc_page(doc, page)
        else:
            add_body_page(doc, page)
        if idx != len(pages) - 1:
            doc.add_page_break()
    doc.save(out_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("out", type=Path)
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)
    build(args.out)
