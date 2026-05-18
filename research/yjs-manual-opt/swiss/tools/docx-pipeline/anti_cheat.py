#!/usr/bin/env python3
"""anti-cheat 三道闸 + 辅助检查 (QA-RULES §8.2)。

骨架来源：`.claude/skills/docx-pipeline/references/anti-cheat-impl.md`
SOT 阈值定义在 `swiss/QA-RULES.md §8.2`，任何阈值变动 → 改 SOT 后同步本脚本。

用法：
  python anti_cheat.py output/imt050-wevac-eu-cn.docx \
      --baseline 00_discussions/.../final/imt050-wevac-eu-cn.docx
  退出码：0 = 全通过，1 = 任一 ERROR
"""

import argparse
import re
import sys
import zipfile
from pathlib import Path


def count_wt(docx_path: Path) -> int:
    """w:t 节点数。阈值 ≥ 300。"""
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return len(re.findall(r"<w:t[ >/]", xml))


def count_media(docx_path: Path) -> int:
    """word/media/ 文件数。比 baseline 多 ≥ 2 视为 image_hack。"""
    with zipfile.ZipFile(docx_path) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def count_chars(docx_path: Path) -> int:
    """所有文本字符数（用于 text_ratio）。"""
    from docx import Document
    doc = Document(str(docx_path))
    total = sum(len(p.text) for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                total += len(cell.text)
    return total


def count_pages(docx_path: Path) -> int:
    """LO headless 转 PDF → 数页数。"""
    import os
    import platform
    import shutil
    import subprocess
    import tempfile
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        if platform.system() == "Windows":
            soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
        elif platform.system() == "Darwin":
            soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        else:
            raise RuntimeError("soffice not found on PATH")
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(
            [soffice, "--headless", "--norestore", "--nolockcheck",
             "--convert-to", "pdf",
             "--outdir", tmp, str(docx_path)],
            check=True, capture_output=True, timeout=180,
        )
        pdf_path = Path(tmp) / (docx_path.stem + ".pdf")
        try:
            import fitz
            return len(fitz.open(pdf_path))
        except ImportError:
            import pypdf
            return len(pypdf.PdfReader(str(pdf_path)).pages)


def check_word_com(docx_path: Path) -> bool:
    """docx2pdf 试转 → 不报错即 Word COM 可打开。W28 教训硬约束。"""
    try:
        from docx2pdf import convert
        import tempfile
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "word_check.pdf"
            convert(str(docx_path), str(out))
            return out.exists() and out.stat().st_size > 0
    except Exception as exc:
        print(f"   word_com exception: {exc}", file=sys.stderr)
        return False


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    ap.add_argument("--baseline", type=Path, required=True,
                    help="W50 cn docx for ratio baselines")
    ap.add_argument("--skip-word-com", action="store_true",
                    help="本机无 MS Word 时跳过（仅 dev）")
    args = ap.parse_args()

    results: list[tuple[str, bool, str]] = []

    # 1. wt_count >= 300
    wt = count_wt(args.docx)
    results.append(("wt_count", wt >= 300, f"wt_count={wt}"))

    # 2. image_hack
    media_actual = count_media(args.docx)
    media_baseline = count_media(args.baseline)
    image_hack = media_actual > media_baseline + 1
    results.append(("image_hack", not image_hack,
                    f"media={media_actual} baseline={media_baseline}"))

    # 3. text_ratio
    # CN baseline char counts include CJK (1 char dense). For non-CN langs the
    # baseline is the SAME-LANG round-trip if available (output/...-{lang}.docx
    # baseline file path can be passed). For CN it stays vs W50 source.
    # Default [0.95, 1.20] only applies when comparing within same language.
    actual = count_chars(args.docx)
    baseline = count_chars(args.baseline)
    ratio = actual / baseline if baseline else 0
    # Heuristic: if baseline is CN and target is non-CN, accept [0.95, 4.0]
    # (CN→EN typically expands 2.5-3.3x because CJK is character-dense).
    target_is_cn = "-cn" in args.docx.name or "/cn." in str(args.docx)
    baseline_is_cn = "-cn" in args.baseline.name
    if baseline_is_cn and not target_is_cn:
        # cross-language ratio
        rng = (0.80, 4.00)
        rng_label = "cross-lang [0.80, 4.00]"
    else:
        rng = (0.95, 1.20)
        rng_label = "[0.95, 1.20]"
    results.append(("text_ratio", rng[0] <= ratio <= rng[1],
                    f"ratio={ratio:.3f} ({actual}/{baseline}) range={rng_label}"))

    # 4. page count = 15 (+-1)
    try:
        pages = count_pages(args.docx)
        results.append(("page_count", 14 <= pages <= 16, f"pages={pages}"))
    except Exception as exc:
        results.append(("page_count", False, f"error: {exc}"))

    # 5. Word COM
    if not args.skip_word_com:
        results.append(("word_com", check_word_com(args.docx), "docx2pdf"))
    else:
        results.append(("word_com", True, "SKIPPED (--skip-word-com)"))

    ok = all(passed for _, passed, _ in results)
    for name, passed, detail in results:
        mark = "PASS" if passed else "FAIL"
        print(f"[{mark}] {name}: {detail}")

    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
